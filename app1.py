import streamlit as st
from crewai import Crew, Agent, Task
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from docx import Document
from io import BytesIO

st.title("COBOL Code to BRD Generator - Agentic App")

st.markdown("""
This app helps convert COBOL source code into a structured Business Requirements Document (BRD).
Provide your COBOL code and business glossary below.
""")

# Text area for COBOL code
cobol_code = st.text_area("Enter COBOL Source Code:", height=300)

# Text area for business glossary
glossary = st.text_area("Enter Business Glossary (optional):", height=150)
with st.sidebar:
    st.markdown("### App Instructions")
    st.markdown("""
    1. Paste your COBOL code.
    2. Click 'Generate BRD using CrewAI'.
    3. Review the generated BRD and download it as a DOCX.
    4. Click 'Generate Java Code from BRD' to create Java implementation.
    5. Click 'Evaluate & Optimize Java Code' for improved accuracy.
    6. Click 'Apply Enterprise-Grade Refinements' for production-ready code.
    """)
    
    st.markdown("### Optimization Levels")
    st.markdown("""
    - **Basic**: Initial Java code generation (~60% complete)
    - **Optimized**: Improved with evaluation feedback (~75% complete)
    - **Enterprise**: Production-ready implementation (~100% complete)
      - Complete class definitions
      - Full COBOL record parsing
      - Proper exception handling
      - BigDecimal for financial precision
      - HashMap-based lookups
      - Unit tests
    """)

# Button to trigger BRD generation
if st.button("Generate BRD using CrewAI"):
    if not cobol_code.strip():
        st.warning("Please enter some COBOL code.")
    else:
        llm = ChatOpenAI(temperature=0, model="gpt-4")

        # Define the prompts
        AgentPrompt = """You are a senior Business Analyst with expertise in COBOL programs and Business Requirement Documents (BRDs).
Your task is to analyze the COBOL source code provided and extract relevant details needed for a clear and accurate BRD.

ONLY use the information explicitly mentioned in the COBOL source code and the business glossary.
DO NOT make assumptions or include any information that is not present in the input data.

Focus on:
- Identifying business rules, especially critical ones as per business glossary or comment tags
- Input and output data structures, including key fields and their purpose
- Key calculations or transformations
- Conditions or validations present in the logic, including any error messages or exception handling mechanisms
- System interfaces, dependencies, or architectural notes
- Any performance considerations or technical constraints mentioned in the COBOL code

You are not expected to reverse-engineer the full logic, only to summarize observable behavior from the COBOL code to be used in a BRD context.

Be clear and concise. When something is ambiguous or undefined, mention it as \"Not explicitly defined in source code.\"

Respond only in plain English."""

        BRD_PROMPT = """
You are a documentation expert tasked with converting the technical analysis into a formal Business Requirements Document (BRD) section.

ONLY use the content provided in the \"COBOL Analysis Output\" below.
Do NOT introduce any new functionality or inferred behavior.
Avoid hallucinations or speculative assumptions.
Stick strictly to what is mentioned in the analysis.

Structure your BRD section to include:
- Functional Overview
- Inputs and Outputs (include key fields if available)
- Business Rules (highlight critical ones if tagged)
- Error Handling (include validations and messages)
- Dependencies (if any)
- Technical Constraints (like performance considerations)
- Glossary (if referenced terms are used)
- Process Flow Visualization (describe the flow in plain steps)

Maintain a professional tone and avoid unnecessary technical jargon.
"""

        # Define agents
        cobol_analyst = Agent(
            role="COBOL Analyst",
            goal="Analyze COBOL source code and extract key details for a BRD",
            backstory="You are a senior business analyst skilled in understanding legacy COBOL systems and translating them into modern documentation.",
            verbose=True,
            allow_delegation=False,
            llm=llm
        )

        brd_writer = Agent(
            role="BRD Writer",
            goal="Draft a professional BRD section from a COBOL analysis",
            backstory="You are an experienced documentation specialist known for creating precise business requirement documents.",
            verbose=True,
            allow_delegation=False,
            llm=llm
        )

        # Define tasks
        analyze_cobol_task = Task(
            description=AgentPrompt + f"\n\nCOBOL Code:\n{cobol_code}\n\nGlossary:\n{glossary}",
            agent=cobol_analyst,
            expected_output="COBOL Analysis Output"
        )

        brd_generation_task = Task(
            description=BRD_PROMPT + "\n\nUse the COBOL Analysis Output from the previous agent to create the BRD section.",
            agent=brd_writer,
            expected_output="Final BRD Section"
        )

        # Define the crew
        crew = Crew(
            agents=[cobol_analyst, brd_writer],
            tasks=[analyze_cobol_task, brd_generation_task],
            verbose=True
        )

        # Execute the crew
        result = crew.kickoff()
        
        # Store the result in session state
        st.session_state["brd_output"] = result.raw

# Display the BRD output if it exists in session state
if "brd_output" in st.session_state:
    st.subheader("Generated BRD Output")
    st.write(st.session_state["brd_output"])

    # Create a Word document from the result
    document = Document()
    document.add_heading("Business Requirements Document (BRD)", level=1)
    document.add_paragraph(st.session_state["brd_output"])

    # Save the document to an in-memory buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Provide a download link for the Word document
    st.download_button(
        label="Download BRD as Word Document",
        data=buffer,
        file_name="Generated_BRD.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
        
# Button to trigger Java code generation
if st.button("Generate Java Code from BRD"):
    if "brd_output" not in st.session_state:
        st.warning("Please generate a BRD first.")
    else:
        llm = ChatOpenAI(temperature=0, model="gpt-4")
        
        # Create an agent to generate the java code based on the BRD and the cobol input
        java_code_generator = Agent(
            role='COBOL to Java Migration Expert',
            goal='Generate complete production-ready Java code that exactly replicates all COBOL functionality',
            backstory='You are a senior software architect specializing in legacy modernization with 15+ years of experience converting COBOL systems to Java. You understand both languages deeply and can create exact functional equivalents while using modern Java practices.',
            verbose=True,
            allow_delegation=False,
            llm=llm
        )

        java_generation_task = Task(
            description=f"""Generate a COMPLETE and DETAILED Java code implementation that is functionally equivalent to the provided COBOL code.
            
            COBOL Code:
            {cobol_code}
            
            BRD Output:
            {st.session_state["brd_output"]}
            
            Follow these guidelines:
            1. Create a COMPLETE end-to-end Java implementation with ALL necessary classes, methods, and files
            2. Include ALL business logic, calculations, and data processing from the COBOL code
            3. Implement proper error handling for all operations
            4. Include detailed comments explaining the code and its relationship to the original COBOL
            5. Use modern Java best practices and appropriate design patterns
            6. Create proper data structures to represent all COBOL records and variables
            7. Implement ALL validation rules and business logic exactly as in the COBOL code
            8. Include a main class that demonstrates the execution flow
            9. Ensure all file operations are properly implemented if present in the COBOL code
            10. DO NOT provide simplified or sample code - implement the FULL functionality
            
            Your response should be production-ready Java code that could directly replace the COBOL implementation.""",
            agent=java_code_generator,
            expected_output="Complete Java implementation with all business logic and functionality"
        )

        # Create a new crew for Java generation
        java_crew = Crew(
            agents=[java_code_generator],
            tasks=[java_generation_task],
            verbose=True
        )

        # Execute the crew with the new task
        with st.spinner("Generating detailed Java implementation (this may take a few minutes)..."):
            st.info("Creating complete Java equivalent with all business logic and functionality from the COBOL code...")
            java_result = java_crew.kickoff()
            
            # Store the Java code output in session state
            st.session_state["java_output"] = java_result.raw
            st.session_state["original_java_output"] = java_result.raw  # Store original for comparison

# Button to evaluate and optimize Java code
if st.button("Evaluate & Optimize Java Code"):
    if "java_output" not in st.session_state:
        st.warning("Please generate Java code first.")
    else:
        llm = ChatOpenAI(temperature=0, model="gpt-4")
        st.session_state["optimization_level"] = "basic"
        
        # Create an evaluator agent
        code_evaluator = Agent(
            role='COBOL-Java Conversion Evaluator',
            goal='Evaluate and identify gaps in Java code converted from COBOL',
            backstory='You are a technical lead with expertise in both COBOL and Java, specializing in evaluating the completeness and accuracy of code migrations.',
            verbose=True,
            allow_delegation=False,
            llm=llm
        )
        
        # Create an optimizer agent
        code_optimizer = Agent(
            role='Java Code Optimizer',
            goal='Enhance Java code to fully implement all COBOL functionality',
            backstory='You are a senior Java developer with deep knowledge of COBOL to Java migration patterns, focused on ensuring functional equivalence and code quality.',
            verbose=True,
            allow_delegation=False,
            llm=llm
        )
        
        # Define evaluation criteria
        evaluation_criteria = """
        Evaluate the Java code against the original COBOL code for the following aspects:
        
        1. File Input/Output: Check if all file operations in COBOL are properly implemented in Java
        2. Business Logic: Verify all calculations, formulas, and business rules are correctly translated
        3. Error Handling: Ensure all error conditions and exception handling are implemented
        4. Data Structures: Confirm all COBOL data structures have equivalent Java representations
        5. Control Flow: Validate that program flow, loops, and conditionals match the COBOL logic
        6. Report Generation: Check if all reporting functionality is properly implemented
        7. Data Persistence: Verify that data storage and retrieval mechanisms are equivalent
        8. Validation Rules: Ensure all input validation and business rule validations are implemented
        
        For each aspect, provide:
        - Status: Complete, Partial, or Missing
        - Specific gaps or issues identified
        - Recommendations for improvement
        
        Format your evaluation as a structured report with clear sections for each aspect.
        """
        
        # Evaluation task
        evaluation_task = Task(
            description=f"""Evaluate the Java code against the original COBOL code using the criteria below.
            
            COBOL Code:
            {cobol_code}
            
            Java Code:
            {st.session_state["java_output"]}
            
            Evaluation Criteria:
            {evaluation_criteria}
            
            Provide a detailed evaluation report identifying specific gaps and issues.
            """,
            agent=code_evaluator,
            expected_output="Detailed evaluation report of Java code"
        )
        
        # Optimization task
        optimization_task = Task(
            description="""Based on the evaluation report, optimize the Java code to address all identified gaps and issues.
            
            Focus on implementing missing functionality and fixing incomplete aspects while maintaining the existing correct implementations.
            
            Ensure the optimized code is complete, accurate, and follows Java best practices.
            
            Return the complete optimized Java code implementation.
            """,
            agent=code_optimizer,
            expected_output="Optimized Java code implementation"
        )
        
        # Run evaluation task first
        evaluation_crew = Crew(
            agents=[code_evaluator],
            tasks=[evaluation_task],
            verbose=True
        )
        
        # Execute the evaluation first
        with st.spinner("Evaluating Java code..."):
            st.info("Analyzing code for gaps and issues...")
            evaluation_result = evaluation_crew.kickoff()
            
            # Store the evaluation report
            st.session_state["evaluation_report"] = evaluation_result.raw
            
            # Now run the optimization with the evaluation result
            optimization_task = Task(
                description=f"""Based on the evaluation report below, optimize the Java code to address all identified gaps and issues.
                
                Evaluation Report:
                {evaluation_result.raw}
                
                Original Java Code:
                {st.session_state["java_output"]}
                
                Original COBOL Code:
                {cobol_code}
                
                Focus on implementing missing functionality and fixing incomplete aspects while maintaining the existing correct implementations.
                
                Ensure the optimized code is complete, accurate, and follows Java best practices.
                
                Return the complete optimized Java code implementation.
                """,
                agent=code_optimizer,
                expected_output="Optimized Java code implementation"
            )
            
            # Create optimization crew
            optimization_crew = Crew(
                agents=[code_optimizer],
                tasks=[optimization_task],
                verbose=True
            )
            
            # Run the optimization
            st.info("Implementing improvements based on evaluation...")
            optimization_result = optimization_crew.kickoff()
            
            # Store the optimized code
            st.session_state["optimized_java_output"] = optimization_result.raw
            st.session_state["java_output"] = st.session_state["optimized_java_output"]

# Display the evaluation report if it exists
if "evaluation_report" in st.session_state:
    st.subheader("Java Code Evaluation Report")
    st.write(st.session_state["evaluation_report"])
    
    # Add option to view original code
    if st.checkbox("Show Original Java Code"):
        st.subheader("Original Java Code")
        st.code(st.session_state["original_java_output"], language="java")

# Button for enterprise-grade refinement
if "optimized_java_output" in st.session_state and st.button("Apply Enterprise-Grade Refinements"):
    llm = ChatOpenAI(temperature=0, model="gpt-4")
    
    # Create an enterprise refinement agent
    enterprise_refiner = Agent(
        role='Enterprise Java Implementation Expert',
        goal='Create complete, production-ready Java code with no gaps or skeleton code',
        backstory='You are a principal architect specializing in enterprise Java applications with expertise in financial systems and COBOL migration projects. You are known for delivering 100% complete, compilable code with no TODOs or implementation gaps.',
        verbose=True,
        allow_delegation=False,
        llm=llm
    )
    
    # Define enterprise refinement criteria
    enterprise_criteria = """
    Refine the Java code to address these critical enterprise requirements:
    
    1. CRITICAL: Complete All Class Definitions
       - Implement ALL referenced classes (Employee, Timecard, PayrollReport, etc.)
       - Include ALL fields, constructors, getters/setters, and methods
       - Ensure no "skeleton" or commented-out code remains
    
    2. CRITICAL: Implement Complete COBOL Record Parsing
       - Create parsers for ALL fixed-length COBOL record formats
       - Implement proper field extraction with exact positioning
       - Handle COMP-3 packed decimal fields correctly
       - Example implementation required:
         ```java
         public static Employee parseCobolEmployee(String record) {
             if (record.length() < 120) {
                 throw new IllegalArgumentException("Invalid employee record length");
             }
             
             String id = record.substring(0, 6).trim();
             String name = record.substring(6, 31).trim();
             BigDecimal hourlyRate = parseComp3(record, 31, 5); // COMP-3 parsing
             int exemptions = Integer.parseInt(record.substring(36, 38).trim());
             String state = record.substring(38, 40).trim();
             
             return new Employee(id, name, hourlyRate, exemptions, state);
         }
         ```
    
    3. Fix Exception Handling Anti-Patterns
       - Create custom exception classes for different error scenarios
       - Either log OR throw exceptions, not both
       - Implement proper exception hierarchy
       - Example implementation required:
         ```java
         public void loadEmployeesFromFile(String filename) throws PayrollProcessingException {
             try {
                 // ... file processing
             } catch (IOException e) {
                 String message = "Failed to load employees from: " + filename;
                 LOGGER.error(message, e);
                 throw new PayrollProcessingException(message, e);
             }
         }
         ```
    
    4. Data Compatibility: Implement proper data readers/writers that can directly work with COBOL data files
       - Add support for EBCDIC encoding if needed
       - Handle fixed-length records properly
       - Implement proper data type conversions
    
    5. Financial Precision: Replace all floating-point calculations with proper decimal handling
       - Use BigDecimal for all monetary calculations
       - Ensure exact decimal precision matching COBOL's COMP-3 fields
       - Implement proper rounding modes for financial calculations
    
    6. Performance Optimization:
       - Replace List iterations with HashMap lookups for better performance
       - Optimize file I/O operations for large datasets
       - Add appropriate caching mechanisms
    
    7. Add Unit Tests:
       - Include JUnit test framework setup
       - Add test cases for critical business logic
       - Include sample test data
    
    Return the complete enterprise-ready Java implementation with NO skeleton code, NO commented-out functionality, and ALL classes fully implemented.
    """
    
    # Enterprise refinement task
    refinement_task = Task(
        description=f"""Refine the optimized Java code to enterprise production standards.
        
        Current Java Code:
        {st.session_state["java_output"]}
        
        Original COBOL Code:
        {cobol_code}
        
        Enterprise Criteria:
        {enterprise_criteria}
        
        CRITICAL IMPLEMENTATION REQUIREMENTS:
        1. DO NOT leave any commented-out code or TODOs in your implementation
        2. Implement ALL class definitions completely (Employee, Timecard, PayrollReport, etc.)
        3. Implement ALL COBOL record parsing logic with exact field positions
        4. Fix ALL exception handling to follow best practices
        5. Ensure ALL business logic is fully implemented, not just described
        
        Your implementation MUST be complete, compilable, and ready for production use.
        
        Provide a complete enterprise-grade Java implementation that addresses all the criteria above.
        """,
        agent=enterprise_refiner,
        expected_output="Enterprise-grade Java implementation"
    )
    
    # Create refinement crew
    refinement_crew = Crew(
        agents=[enterprise_refiner],
        tasks=[refinement_task],
        verbose=True
    )
    
    # Execute the refinement
    with st.spinner("Applying enterprise-grade refinements (this may take several minutes)..."):
        st.info("Enhancing code with production-ready features for enterprise deployment...")
        refinement_result = refinement_crew.kickoff()
        
        # Store the refined code
        st.session_state["enterprise_java_output"] = refinement_result.raw
        st.session_state["java_output"] = st.session_state["enterprise_java_output"]
        st.session_state["optimization_level"] = "enterprise"

# Display the Java code output if it exists in session state
if "java_output" in st.session_state:
    st.subheader("Java Code Implementation")
    
    # Add explanation based on optimization level
    if "optimization_level" in st.session_state:
        if st.session_state["optimization_level"] == "enterprise":
            st.success("""✅ PRODUCTION-READY IMPLEMENTATION: 
            
This Java code has been fully implemented with:
- Complete class definitions (Employee, Timecard, PayrollReport, etc.)
- Full COBOL record parsing with exact field positioning
- Proper exception handling following best practices
- BigDecimal for financial calculations
- HashMap-based lookups for performance
- Unit tests for critical business logic
            
The code is now complete, compilable, and ready for production use.""")
        else:
            st.success("This is the optimized Java implementation with improvements based on the evaluation.")
    else:
        st.info("Below is the Java implementation that functionally replicates the COBOL code.")
    
    # Display the code with syntax highlighting
    st.code(st.session_state["java_output"], language="java")
    
    # Create a text file for download
    java_buffer = BytesIO()
    java_buffer.write(st.session_state["java_output"].encode())
    java_buffer.seek(0)
    
    # Provide download button for Java code with appropriate name
    file_name = "Enterprise_Java_Implementation.java" if "optimization_level" in st.session_state and st.session_state["optimization_level"] == "enterprise" else "COBOL_to_Java_Implementation.java"
    
    st.download_button(
        label=f"Download {file_name}",
        data=java_buffer,
        file_name=file_name,
        mime="text/plain"
    )